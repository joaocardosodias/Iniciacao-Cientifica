"""
generate_test_files.py — Gerador de ambiente de teste.

Cria uma estrutura realista de pastas com milhares de arquivos falsos
(.xlsx, .docx, .pdf, .txt) com conteúdo plausível para simular a
máquina de uma vítima em ambiente de VM isolada.

Uso:
    python generate_test_files.py                   # usa ~/Documentos por padrão
    python generate_test_files.py /caminho/destino  # pasta personalizada
    python generate_test_files.py --count 5000      # quantidade de arquivos

Dependências:
    pip install openpyxl python-docx reportlab faker
"""

import argparse
import os
import random
import sys
from pathlib import Path

# ── Verificação de dependências ────────────────────────────────────────────────

def _check_deps():
    missing = []
    for pkg in ("openpyxl", "docx", "reportlab", "faker"):
        try:
            __import__(pkg)
        except ImportError:
            missing.append(pkg if pkg != "docx" else "python-docx")
    if missing:
        print(f"[ERRO] Dependências faltando: {', '.join(missing)}")
        print(f"       Instale com: pip install {' '.join(missing)}")
        sys.exit(1)

_check_deps()

import openpyxl
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from faker import Faker

fake = Faker("pt_BR")

# ── Estrutura de pastas realistas ──────────────────────────────────────────────

FOLDER_TREE = [
    "Documentos/Financeiro/2023",
    "Documentos/Financeiro/2024",
    "Documentos/Financeiro/2025",
    "Documentos/Financeiro/Notas Fiscais",
    "Documentos/Financeiro/Extratos Bancários",
    "Documentos/RH/Contratos",
    "Documentos/RH/Folha de Pagamento",
    "Documentos/RH/Currículos",
    "Documentos/Jurídico/Contratos Clientes",
    "Documentos/Jurídico/Contratos Fornecedores",
    "Documentos/Jurídico/LGPD",
    "Documentos/Projetos/Alpha",
    "Documentos/Projetos/Beta",
    "Documentos/Projetos/Confidencial",
    "Documentos/Clientes/Ativos",
    "Documentos/Clientes/Inativos",
    "Desktop/Relatórios",
    "Desktop/Reuniões",
    "Downloads/Recebidos",
    "Pictures/Capturas",
]

# Prefixos de nome de arquivo por categoria
FILE_NAMES = {
    "financeiro": [
        "Balancete", "DRE", "Fluxo_de_Caixa", "Orcamento", "Fatura",
        "Nota_Fiscal", "Extrato", "Recibo", "Planilha_Custos", "Budget",
    ],
    "rh": [
        "Contrato_CLT", "Folha_Pagamento", "Ficha_Cadastral", "Ferias",
        "Rescisao", "Admissao", "Curriculo", "Avaliacao_Desempenho",
    ],
    "juridico": [
        "Contrato_Prestacao", "Acordo_NDA", "Termo_Confidencialidade",
        "Procuracao", "Ata_Reuniao", "Politica_Privacidade", "LGPD_Conformidade",
    ],
    "projetos": [
        "Escopo_Projeto", "Cronograma", "Requisitos", "Arquitetura",
        "Proposta_Tecnica", "Ata_Sprint", "Roadmap", "Kickoff",
    ],
    "geral": [
        "Relatorio", "Apresentacao", "Resumo", "Documento", "Arquivo",
        "Notas", "Backup_Dados", "Planilha", "Memorando", "Oficio",
    ],
}

EXTENSIONS = [".xlsx", ".docx", ".pdf", ".txt"]

# Pesos para cada extensão (mais xlsx e docx, como numa empresa real)
EXT_WEIGHTS = [0.35, 0.30, 0.20, 0.15]

# ── Geradores de conteúdo falso ────────────────────────────────────────────────

def _make_txt(path: Path):
    paragraphs = [fake.paragraph(nb_sentences=random.randint(3, 8)) for _ in range(random.randint(2, 6))]
    header = f"{fake.company()} — {fake.catch_phrase()}\nData: {fake.date()}\nAutor: {fake.name()}\n\n"
    path.write_text(header + "\n\n".join(paragraphs), encoding="utf-8")


def _make_xlsx(path: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = random.choice(["Dados", "Relatório", "Planilha", "Resumo"])

    # Cabeçalho
    headers = random.choice([
        ["Nome", "CPF", "Salário", "Departamento", "Data Admissão"],
        ["Produto", "Qtd", "Valor Unit.", "Total", "Fornecedor"],
        ["Cliente", "CNPJ", "Valor Contrato", "Vencimento", "Status"],
        ["Mês", "Receita", "Despesa", "Lucro", "Margem %"],
    ])
    ws.append(headers)

    # Linhas de dados
    for _ in range(random.randint(20, 150)):
        row = []
        for h in headers:
            h_lower = h.lower()
            if "nome" in h_lower or "cliente" in h_lower:
                row.append(fake.name())
            elif "cpf" in h_lower:
                row.append(fake.cpf())
            elif "cnpj" in h_lower:
                row.append(fake.cnpj())
            elif "salário" in h_lower or "valor" in h_lower or "receita" in h_lower or "despesa" in h_lower:
                row.append(round(random.uniform(1500, 150000), 2))
            elif "data" in h_lower or "vencimento" in h_lower or "mês" in h_lower:
                row.append(str(fake.date()))
            elif "departamento" in h_lower or "status" in h_lower or "fornecedor" in h_lower:
                row.append(fake.company())
            elif "%" in h_lower or "margem" in h_lower:
                row.append(round(random.uniform(1, 45), 2))
            else:
                row.append(random.randint(1, 999))
        ws.append(row)

    wb.save(path)


def _make_docx(path: Path):
    doc = Document()
    doc.add_heading(fake.catch_phrase(), 0)
    doc.add_paragraph(f"Empresa: {fake.company()}  |  Data: {fake.date()}  |  Responsável: {fake.name()}")
    doc.add_paragraph("")

    for _ in range(random.randint(2, 5)):
        doc.add_heading(fake.bs().title(), level=1)
        for _ in range(random.randint(1, 4)):
            doc.add_paragraph(fake.paragraph(nb_sentences=random.randint(3, 7)))

    doc.save(path)


def _make_pdf(path: Path):
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4

    # Título
    c.setFont("Helvetica-Bold", 16)
    title = fake.catch_phrase()[:60]
    c.drawString(50, height - 60, title)

    # Subtítulo
    c.setFont("Helvetica", 10)
    c.drawString(50, height - 80, f"{fake.company()} | {fake.date()} | {fake.name()}")

    # Corpo
    c.setFont("Helvetica", 10)
    y = height - 120
    for _ in range(random.randint(3, 7)):
        # Parágrafo simulado (quebra em linhas de 90 chars)
        paragraph = fake.paragraph(nb_sentences=random.randint(3, 6))
        words = paragraph.split()
        line = ""
        for word in words:
            if len(line) + len(word) + 1 > 90:
                c.drawString(50, y, line)
                y -= 15
                line = word
                if y < 60:
                    c.showPage()
                    y = height - 60
                    c.setFont("Helvetica", 10)
            else:
                line = f"{line} {word}".strip()
        if line:
            c.drawString(50, y, line)
            y -= 25

    c.save()


_GENERATORS = {
    ".txt":  _make_txt,
    ".xlsx": _make_xlsx,
    ".docx": _make_docx,
    ".pdf":  _make_pdf,
}

# ── Lógica principal ───────────────────────────────────────────────────────────

def _random_filename(category: str | None = None) -> str:
    cat = category or random.choice(list(FILE_NAMES.keys()))
    prefix = random.choice(FILE_NAMES[cat])
    suffix = random.choice([
        fake.date().replace("-", ""),
        str(random.randint(100, 9999)),
        fake.last_name().replace(" ", "_"),
        f"v{random.randint(1,5)}",
        f"rev{random.randint(1,3)}",
    ])
    return f"{prefix}_{suffix}"


def generate(base_dir: Path, total: int):
    """
    Gera `total` arquivos falsos distribuídos pela estrutura de pastas.

    Args:
        base_dir: Pasta raiz onde os arquivos serão criados.
        total:    Número total de arquivos a gerar.
    """
    # Cria todas as pastas
    folders = []
    for rel in FOLDER_TREE:
        folder = base_dir / rel
        folder.mkdir(parents=True, exist_ok=True)
        folders.append(folder)

    print(f"\n[+] Destino  : {base_dir}")
    print(f"[+] Pastas   : {len(folders)}")
    print(f"[+] Arquivos : {total}")
    print(f"[+] Tipos    : .xlsx .docx .pdf .txt\n")

    errors = 0
    for i in range(1, total + 1):
        folder = random.choice(folders)
        ext    = random.choices(EXTENSIONS, weights=EXT_WEIGHTS)[0]
        name   = _random_filename()
        path   = folder / f"{name}{ext}"

        # Evita sobrescrever: adiciona número se já existir
        counter = 1
        while path.exists():
            path = folder / f"{name}_{counter}{ext}"
            counter += 1

        try:
            _GENERATORS[ext](path)
        except Exception as e:
            errors += 1
            print(f"  [WARN] Falha ao criar {path.name}: {e}")
            continue

        # Progresso a cada 100 arquivos
        if i % 100 == 0 or i == total:
            bar_len = 30
            filled  = int(bar_len * i / total)
            bar     = "█" * filled + "░" * (bar_len - filled)
            pct     = i / total * 100
            print(f"\r  [{bar}] {pct:5.1f}%  ({i}/{total})", end="", flush=True)

    print(f"\n\n[✓] Concluído! {total - errors} arquivo(s) criados, {errors} erro(s).\n")


# ── Entry point ────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Gera arquivos falsos (.xlsx, .docx, .pdf, .txt) para ambiente de teste."
    )
    parser.add_argument(
        "dest",
        nargs="?",
        default=str(Path.home() / "Documentos_Teste"),
        help="Pasta de destino (padrão: ~/Documentos_Teste)",
    )
    parser.add_argument(
        "--count", "-n",
        type=int,
        default=5000,
        help="Número total de arquivos a gerar (padrão: 2000)",
    )
    args = parser.parse_args()

    base_dir = Path(args.dest).expanduser().resolve()
    generate(base_dir, args.count)


if __name__ == "__main__":
    main()
